#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 文档生成工具 —— 龙虾优盘「通用工具」能力型技能

纯对话模型只能给文字;本脚本把结构化内容(JSON 规格)渲染成**真实的 .docx 文件**,
带标题层级、项目符号、编号列表、表格、分页,并配置中文默认字体。

用法:
  python docx_tool.py build --spec 规格.json -o 文档.docx
  python docx_tool.py build --spec -        -o 文档.docx   # 从标准输入读规格

规格(JSON):
  {
    "title": "文档标题", "subtitle": "可选副标题",
    "blocks": [
      {"type":"h1","text":"一级标题"},
      {"type":"h2","text":"二级标题"},
      {"type":"p","text":"正文段落"},
      {"type":"bullets","items":["要点一","要点二"]},
      {"type":"numbered","items":["第一步","第二步"]},
      {"type":"table","header":["列A","列B"],"rows":[["1","2"]]},
      {"type":"pagebreak"}
    ],
    "footer": "可选页脚"
  }

依赖: python-docx>=1.0。退出码: 0 成功; 1 参数/依赖/规格错误; 2 处理错误。
"""
import argparse
import json
import os
import sys

try:
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
    sys.stdin.reconfigure(encoding="utf-8")
except Exception:
    pass

# 常见中文字体候选(按系统可用性择一),保证中文正常显示
_CJK_FONTS = ["微软雅黑", "Microsoft YaHei", "宋体", "SimSun", "PingFang SC"]


def _require_docx():
    try:
        import docx  # noqa: F401
        return docx
    except ImportError:
        sys.stderr.write('[docx_tool] 缺少依赖 python-docx,请先执行: pip install "python-docx>=1.0"\n')
        sys.exit(1)


def _set_cjk_font(doc, font_name: str):
    """把正文默认字体设为中文字体(含 eastAsia),避免中文用西文字体渲染。"""
    from docx.oxml.ns import qn
    style = doc.styles["Normal"]
    style.font.name = font_name
    try:
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass


def _pick_cjk_font() -> str:
    """从候选里选一个(简单返回首个;Word 打开时缺字体会自行回退)。"""
    return _CJK_FONTS[0]


def _add_table(doc, block):
    header = block.get("header") or []
    rows = block.get("rows") or []
    n_cols = max(len(header), max((len(r) for r in rows), default=0))
    if n_cols == 0:
        return
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    if header:
        for j, cell_text in enumerate(header):
            table.rows[0].cells[j].text = str(cell_text)
            for run in table.rows[0].cells[j].paragraphs[0].runs:
                run.bold = True
    for r in rows:
        cells = table.add_row().cells
        for j in range(n_cols):
            cells[j].text = str(r[j]) if j < len(r) else ""


def build(spec: dict, output: str):
    docx = _require_docx()
    from docx.shared import Pt

    if not isinstance(spec, dict):
        sys.stderr.write("[docx_tool] 规格必须是 JSON 对象\n")
        sys.exit(1)
    blocks = spec.get("blocks") or []

    doc = docx.Document()
    _set_cjk_font(doc, _pick_cjk_font())

    if spec.get("title"):
        t = doc.add_heading(str(spec["title"]), level=0)
        t.alignment = 1  # 居中
    if spec.get("subtitle"):
        sub = doc.add_paragraph(str(spec["subtitle"]))
        sub.alignment = 1
        for run in sub.runs:
            run.italic = True

    for i, block in enumerate(blocks):
        if not isinstance(block, dict):
            continue
        btype = block.get("type")
        text = str(block.get("text", ""))
        if btype == "h1":
            doc.add_heading(text, level=1)
        elif btype == "h2":
            doc.add_heading(text, level=2)
        elif btype == "h3":
            doc.add_heading(text, level=3)
        elif btype in ("p", "paragraph"):
            doc.add_paragraph(text)
        elif btype == "bullets":
            for item in block.get("items") or []:
                doc.add_paragraph(str(item), style="List Bullet")
        elif btype == "numbered":
            for item in block.get("items") or []:
                doc.add_paragraph(str(item), style="List Number")
        elif btype == "table":
            _add_table(doc, block)
        elif btype == "pagebreak":
            doc.add_page_break()
        else:
            sys.stderr.write(f"[docx_tool] 第 {i+1} 个 block 类型未知: {btype}(已跳过)\n")

    if spec.get("footer"):
        fp = doc.add_paragraph()
        run = fp.add_run(str(spec["footer"]))
        run.font.size = Pt(9)
        run.italic = True

    try:
        doc.save(output)
    except Exception as e:  # noqa: BLE001
        sys.stderr.write(f"[docx_tool] 保存失败: {e}\n")
        sys.exit(2)
    size = os.path.getsize(output)
    print(f"已生成 Word 文档 -> {output}(blocks {len(blocks)} 个, {size} 字节)")


def cmd_build(args):
    if args.spec == "-":
        raw = sys.stdin.read()
    else:
        if not os.path.isfile(args.spec):
            sys.stderr.write(f"[docx_tool] 规格文件不存在: {args.spec}\n")
            sys.exit(1)
        with open(args.spec, "r", encoding="utf-8") as f:
            raw = f.read()
    try:
        spec = json.loads(raw)
    except json.JSONDecodeError as e:
        sys.stderr.write(f"[docx_tool] 规格 JSON 解析失败: {e}\n")
        sys.exit(1)
    build(spec, args.output)


def main():
    parser = argparse.ArgumentParser(prog="docx_tool.py", description="Word 文档生成工具(龙虾优盘通用技能)")
    sub = parser.add_subparsers(dest="command", required=True)
    p = sub.add_parser("build", help="按 JSON 规格生成 .docx")
    p.add_argument("--spec", required=True, help="规格 JSON 文件,或 - 表示标准输入")
    p.add_argument("-o", "--output", required=True, help="输出 .docx 路径")
    p.set_defaults(func=cmd_build)
    args = parser.parse_args()
    args.func(args)


if __name__ == "__main__":
    main()

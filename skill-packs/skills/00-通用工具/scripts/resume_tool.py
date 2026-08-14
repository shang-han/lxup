#!/usr/bin/env python3
# -*- coding: utf-8 -*-
try:
    import sys
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass

import argparse, json, sys
try:
    import docx
except ImportError:
    sys.stderr.write('[resume_tool] 缺少依赖 python-docx,请先执行: pip install "python-docx>=1.0"\n'); sys.exit(1)

def main():
    p = argparse.ArgumentParser()
    p.add_argument('create'); p.add_argument('-o','--output', required=True); p.add_argument('--data', required=True)
    a = p.parse_args()
    try:
        d = json.load(open(a.data, encoding='utf-8'))
    except Exception as e:
        sys.stderr.write('[resume_tool] 无法读取 data: ' + str(e) + '\n'); sys.exit(1)
    missing = [k for k in ('name','phone','email') if not d.get(k)]
    doc = docx.Document()
    doc.add_heading(d.get('name',''), level=0)
    line = ' | '.join(x for x in [d.get('phone',''), d.get('email','')] if x)
    if d.get('target'): line += '  |  求职意向: ' + d['target']
    doc.add_paragraph(line)
    if d.get('summary'):
        doc.add_heading('个人简介', level=1); doc.add_paragraph(d['summary'])
    if d.get('experience'):
        doc.add_heading('工作经历', level=1)
        for e in d['experience']:
            para = doc.add_paragraph(e.get('company','') + ' - ' + e.get('title','') + ' (' + e.get('period','') + ')')
            para.runs[0].bold = True
            for it in e.get('items', []): doc.add_paragraph(it, style='List Bullet')
    if d.get('education'):
        doc.add_heading('教育经历', level=1)
        for e in d['education']:
            doc.add_paragraph(e.get('school','') + ' / ' + e.get('major','') + ' / ' + e.get('degree','') + ' (' + e.get('period','') + ')')
    if d.get('skills'):
        doc.add_heading('技能', level=1); doc.add_paragraph('、'.join(d['skills']))
    doc.save(a.output)
    msg = '已生成 ' + a.output
    if missing: msg += '。缺失字段: ' + ', '.join(missing)
    print(msg)

main()

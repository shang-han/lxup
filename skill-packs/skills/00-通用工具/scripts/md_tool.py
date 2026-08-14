#!/usr/bin/env python3
# -*- coding: utf-8 -*-
try:
    import sys
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass

import argparse, re, sys
try:
    import markdown
except ImportError:
    sys.stderr.write('[md_tool] 缺少依赖 markdown,请先执行: pip install "markdown"\n'); sys.exit(1)

HTML_TPL = '<!DOCTYPE html><html lang="zh-CN"><head><meta charset="utf-8">\n<meta name="viewport" content="width=device-width,initial-scale=1">\n<style>body{max-width:820px;margin:32px auto;padding:0 20px;font:15px/1.8 -apple-system,"Microsoft YaHei",sans-serif;color:#222}\nh1{border-bottom:2px solid #c62828;padding-bottom:8px}h2{border-bottom:1px solid #ddd;padding-bottom:6px}\ncode{background:#f4f4f4;padding:2px 6px;border-radius:4px;font-size:13px}pre{background:#f4f4f4;padding:12px;border-radius:6px;overflow:auto}\ntable{border-collapse:collapse;width:100%}td,th{border:1px solid #ddd;padding:6px 10px}th{background:#fafafa}</style>\n<title>{title}</title></head><body>{body}</body></html>'

def md_to_docx(text, out):
    try:
        import docx
    except ImportError:
        sys.stderr.write('[md_tool] 转 docx 需 python-docx: pip install "python-docx>=1.0"\n'); sys.exit(1)
    d = docx.Document()
    for ln in text.splitlines():
        s = ln.rstrip()
        m = re.match(r'^(#{1,6})\s+(.*)', s)
        if m:
            d.add_heading(m.group(2), level=len(m.group(1)))
        elif re.match(r'^```', s):
            continue
        elif s.startswith('- '):
            d.add_paragraph(s[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s+', s):
            d.add_paragraph(re.sub(r'^\d+\.\s+','',s), style='List Number')
        elif s.strip():
            d.add_paragraph(s)
    d.save(out)

def convert_one(src, out):
    text = open(src, encoding='utf-8').read()
    if out.lower().endswith('.html'):
        title = src.rsplit('/',1)[-1].rsplit('\\',1)[-1]
        open(out, 'w', encoding='utf-8').write(HTML_TPL.format(title=title, body=markdown.markdown(text, extensions=['extra','tables'])))
    elif out.lower().endswith('.docx'):
        md_to_docx(text, out)
    else:
        sys.stderr.write('[md_tool] 输出扩展名需为 .docx 或 .html\n'); sys.exit(1)

def main():
    p = argparse.ArgumentParser()
    p.add_argument('convert'); p.add_argument('src'); p.add_argument('-o','--output'); p.add_argument('--format', choices=['docx','html'])
    a = p.parse_args()
    src = a.src
    import pathlib
    if src.lower().endswith('.md'):
        out = a.output or src[:-3] + '.docx'
        convert_one(src, out); print('已转换 ' + src + ' -> ' + out)
    else:
        d = pathlib.Path(src)
        if not d.is_dir():
            sys.stderr.write('[md_tool] 输入需为 .md 文件或目录\n'); sys.exit(1)
        outdir = pathlib.Path(a.output) if a.output else d
        outdir.mkdir(parents=True, exist_ok=True)
        fmt = a.format or 'docx'
        files = sorted(d.rglob('*.md'))
        for f in files:
            o = outdir / (f.stem + ('.docx' if fmt=='docx' else '.html'))
            convert_one(str(f), str(o))
        print('已转换 ' + str(len(files)) + ' 个文件到 ' + str(outdir))

main()
